#!/usr/bin/env python3
"""
旅游行程文档解析器
支持从PDF和Word文档中提取文本内容
"""

import sys
import argparse
from pathlib import Path

try:
    import PyPDF2
    from docx import Document
except ImportError as e:
    print(f"错误：缺少必要的依赖包。请执行：pip install PyPDF2 python-docx")
    sys.exit(1)


def extract_text_from_pdf(pdf_path):
    """
    从PDF文件中提取文本内容

    Args:
        pdf_path: PDF文件路径

    Returns:
        str: 提取的完整文本内容
    """
    text_content = []

    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # 遍历所有页面
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()

                if text.strip():
                    text_content.append(text)

    except FileNotFoundError:
        raise FileNotFoundError(f"PDF文件不存在：{pdf_path}")
    except Exception as e:
        raise Exception(f"PDF解析失败：{str(e)}")

    return '\n\n'.join(text_content)


def extract_text_from_word(docx_path):
    """
    从Word文档中提取文本内容

    Args:
        docx_path: Word文档路径

    Returns:
        str: 提取的完整文本内容
    """
    text_content = []

    try:
        doc = Document(docx_path)

        # 提取所有段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))

    except FileNotFoundError:
        raise FileNotFoundError(f"Word文件不存在：{docx_path}")
    except Exception as e:
        raise Exception(f"Word文档解析失败：{str(e)}")

    return '\n\n'.join(text_content)


def parse_document(file_path):
    """
    解析文档文件（PDF或Word）并提取文本内容

    Args:
        file_path: 文档文件路径

    Returns:
        str: 提取的完整文本内容

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
        Exception: 解析失败
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{file_path}")

    file_ext = path.suffix.lower()

    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return extract_text_from_word(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{file_ext}。仅支持PDF和Word文档")


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(
        description='旅游行程文档解析器',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    parser.add_argument(
        '--file-path',
        type=str,
        required=True,
        help='文档文件路径（PDF或Word格式）'
    )

    args = parser.parse_args()

    try:
        # 解析文档
        text = parse_document(args.file_path)

        # 输出提取的文本内容
        print(text)

        return 0

    except Exception as e:
        print(f"错误：{str(e)}", file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
